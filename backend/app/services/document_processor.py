import os
import json
import csv
from typing import Optional
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
from openpyxl import load_workbook


class DocumentProcessor:
    """Service for extracting text from PDF, DOCX, Excel, JSON, and CSV documents."""

    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract text from a document based on file type."""
        extension = os.path.splitext(file_path)[1].lower()

        if extension == ".pdf":
            return DocumentProcessor._extract_from_pdf(file_path)
        elif extension == ".docx":
            return DocumentProcessor._extract_from_docx(file_path)
        elif extension in [".xlsx", ".xls"]:
            return DocumentProcessor._extract_from_excel(file_path)
        elif extension == ".json":
            return DocumentProcessor._extract_from_json(file_path)
        elif extension == ".csv":
            return DocumentProcessor._extract_from_csv(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")

    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF using pdfplumber (better layout preservation)."""
        text_parts = []

        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {page_num}]\n{page_text}")
        except Exception as e:
            # Fallback to PyPDF2 if pdfplumber fails
            text_parts = DocumentProcessor._extract_from_pdf_fallback(file_path)

        return "\n\n".join(text_parts)

    @staticmethod
    def _extract_from_pdf_fallback(file_path: str) -> list[str]:
        """Fallback PDF extraction using PyPDF2."""
        text_parts = []

        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {page_num}]\n{page_text}")

        return text_parts

    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file."""
        doc = DocxDocument(file_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    @staticmethod
    def _extract_from_excel(file_path: str) -> str:
        """Extract text from Excel file."""
        wb = load_workbook(file_path, data_only=True)
        text_parts = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"[Sheet: {sheet_name}]")

            # Get headers from first row
            headers = []
            for cell in sheet[1]:
                headers.append(str(cell.value) if cell.value else "")

            if any(headers):
                text_parts.append("Headers: " + " | ".join(headers))

            # Extract data rows
            for row_num, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), 2):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                if any(row_values):
                    # Create key-value pairs using headers
                    if headers:
                        pairs = [f"{h}: {v}" for h, v in zip(headers, row_values) if v]
                        text_parts.append(f"Row {row_num}: " + ", ".join(pairs))
                    else:
                        text_parts.append(" | ".join(row_values))

        return "\n\n".join(text_parts)

    @staticmethod
    def _extract_from_json(file_path: str) -> str:
        """Extract text from JSON file."""
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        def flatten_json(obj, prefix=""):
            """Recursively flatten JSON to readable text."""
            lines = []
            if isinstance(obj, dict):
                for key, value in obj.items():
                    new_prefix = f"{prefix}.{key}" if prefix else key
                    lines.extend(flatten_json(value, new_prefix))
            elif isinstance(obj, list):
                for i, item in enumerate(obj):
                    new_prefix = f"{prefix}[{i}]"
                    lines.extend(flatten_json(item, new_prefix))
            else:
                lines.append(f"{prefix}: {obj}")
            return lines

        text_lines = flatten_json(data)
        return "\n".join(text_lines)

    @staticmethod
    def _extract_from_csv(file_path: str) -> str:
        """Extract text from CSV file."""
        text_parts = []

        # Try to detect encoding
        encodings = ["utf-8", "latin-1", "cp1252"]

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding, newline="") as f:
                    reader = csv.reader(f)
                    rows = list(reader)
                    break
            except UnicodeDecodeError:
                continue
        else:
            raise ValueError("Could not decode CSV file with supported encodings")

        if not rows:
            return ""

        # Use first row as headers
        headers = rows[0] if rows else []
        text_parts.append("Headers: " + " | ".join(headers))

        # Extract data rows
        for row_num, row in enumerate(rows[1:], 2):
            if any(row):
                if headers:
                    pairs = [f"{h}: {v}" for h, v in zip(headers, row) if v]
                    text_parts.append(f"Row {row_num}: " + ", ".join(pairs))
                else:
                    text_parts.append(" | ".join(row))

        return "\n\n".join(text_parts)

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 8000, overlap: int = 500) -> list[str]:
        """Split text into overlapping chunks for processing large documents."""
        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            # Try to find a natural break point (paragraph or sentence)
            if end < len(text):
                # Look for paragraph break
                break_point = text.rfind("\n\n", start + chunk_size - overlap, end)
                if break_point == -1:
                    # Look for sentence break
                    break_point = text.rfind(". ", start + chunk_size - overlap, end)
                if break_point != -1:
                    end = break_point + 2

            chunks.append(text[start:end].strip())
            start = end - overlap

        return chunks

    @staticmethod
    def get_document_stats(text: str) -> dict:
        """Get basic statistics about the document."""
        words = text.split()
        sentences = text.count(".") + text.count("!") + text.count("?")
        paragraphs = text.count("\n\n") + 1

        return {
            "character_count": len(text),
            "word_count": len(words),
            "sentence_count": sentences,
            "paragraph_count": paragraphs,
        }
